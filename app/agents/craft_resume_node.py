from typing import Dict, Any, List
import os
import json
from langchain_core.runnables import RunnableConfig
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage
from google.cloud import secretmanager
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.types.node_state import NodeState
from app.config import settings


class CraftResumeNode:
    """
    A node that creates a professional resume from user-provided data, saves it as a Word document,
    uploads it to Google Drive, and generates a shareable link.
    """

    def __init__(self):
        """
        Initializes the CraftResumeNode by authenticating with Google Cloud and
        setting up Google Docs and Drive API clients.
        """
        # Authenticate with Google Cloud
        credentials = self.authenticate()
        self.doc_client = build("docs", "v1", credentials=credentials)
        self.drive_client = build("drive", "v3", credentials=credentials)

    def create_resume_from_json(self, data: Dict[str, Any]) -> Document:
        """
        Generates a resume as a Word document using the data provided in JSON format.

        Args:
            data (Dict[str, Any]): User's resume data, including personal details, education,
                                   work experience, projects, and skills.

        Returns:
            Document: A Word document object representing the resume.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

        # Helper function to set page margins and size
        def set_page_setup(section):
            # Set Letter size: 8.5" x 11"
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            # Set margins: Top 0.36", Bottom 0.36", Left 0.5", Right 0.5"
            section.top_margin = Inches(0.36)
            section.bottom_margin = Inches(0.36)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Helper function to add a heading
        def add_heading(doc, text, size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, right_text=None):
            paragraph = doc.add_paragraph()
            if right_text:
                # Add tab stops for right alignment
                tab_stops = paragraph.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            run = paragraph.add_run(text)
            run.font.size = Pt(size)
            run.bold = bold
            paragraph.alignment = alignment
            if right_text:
                paragraph.add_run("\t" + right_text)  # Add right-aligned text
            paragraph.paragraph_format.line_spacing = Pt(12)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)

        # Helper function to add a bullet point
        def add_bullet(doc, text):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(text)
            paragraph.paragraph_format.line_spacing = Pt(12)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)

        # Helper function to add a paragraph
        def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
            paragraph = doc.add_paragraph(text)
            paragraph.alignment = alignment
            paragraph.paragraph_format.line_spacing = Pt(12)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)

        # Helper function to add a horizontal line
        def add_horizontal_line(doc):
            # Add a simple horizontal rule using underscores
            paragraph = doc.add_paragraph("_" * 120)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Start document
        doc = Document()
        set_page_setup(doc.sections[0])  # Apply page setup

        # Add personal details
        personal = data["personal_details"]
        add_heading(doc, personal["name"], size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)  # Center-aligned name
        contact_info = f'{personal["phone"]} | {personal["email"]} | {personal["linkedin"]} | {personal["github"]}'
        add_paragraph(doc, contact_info, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Add Education section
        add_horizontal_line(doc)  # Add horizontal line before each major section
        add_heading(doc, "\nEducation", size=14)
        for edu in data["education"]:
            edu_title = f'{edu["degree"]}, {edu["institution"]}'
            edu_dates = f'{edu["start_date"]} - {edu["end_date"]}'
            add_heading(doc, edu_title, size=12, bold=True, right_text=edu_dates)
            add_paragraph(doc, f"GPA: {edu['gpa']}")
            courses = ", ".join(edu["courses"])
            add_paragraph(doc, f"Courses: {courses}")

        # Add Work Experience section
        add_horizontal_line(doc)
        add_heading(doc, "\nWork Experience", size=14)
        for exp in data["experiences"]:
            exp_title = f'{exp["position"]} at {exp["organization"]}'
            exp_dates = f'{exp["start_date"]} - {exp["end_date"]}'
            add_heading(doc, exp_title, size=12, bold=True, right_text=exp_dates)
            for desc in exp["description"]:
                add_bullet(doc, desc)

        # Add Projects section
        add_horizontal_line(doc)
        add_heading(doc, "\nProjects", size=14)
        for proj in data["projects"]:
            proj_title = proj["name"]
            proj_dates = f'{proj["start_date"]} - {proj["end_date"]}'
            add_heading(doc, proj_title, size=12, bold=True, right_text=proj_dates)
            add_paragraph(doc, proj["github_link"])
            for desc in proj["description"]:
                add_bullet(doc, desc)

        # Add Skills section
        add_horizontal_line(doc)
        add_heading(doc, "\nSkills", size=14)
        skills_text = ", ".join(data["skills"])
        add_paragraph(doc, skills_text)

        return doc

    def __call__(self, state: NodeState, config: RunnableConfig) -> Dict[str, Any] | None:
        """
        Orchestrates the process of creating a resume, uploading it to Google Drive,
        and returning a sharable link.

        Args:
            state (NodeState): The input state containing user data.
            config (RunnableConfig): The runtime configuration for the node.

        Returns:
            Dict[str, Any] | None: A response message containing the sharable link.
        """
        # Create the resume
        resume_doc = self.create_resume_from_json(state.user_details)

        # Save the document to a local file
        file_path = f"/tmp/{state.user_details['personal_details']['name']}_resume.docx"
        resume_doc.save(file_path)

        # Upload the file to Google Docs
        resume_link = self.upload_doc_to_google(file_path, state.user_details["personal_details"]["name"])

        # Clean up the local file
        os.remove(file_path)

        return {
            "messages": [
                AIMessage(content=f"I have crafted a draft resume. Here's the link:\n{resume_link}")
            ]
        }

    def upload_doc_to_google(self, file_path: str, name: str) -> str:
        """
        Uploads a .docx file to Google Drive and converts it to Google Docs format.

        Args:
            file_path (str): The path of the file to upload.
            name (str): The name of the user, used to name the file in Drive.

        Returns:
            str: A sharable link to the uploaded Google Doc.
        """
        # Metadata for the file
        file_metadata = {"name": f"{name} - Resume", "mimeType": "application/vnd.google-apps.document"}
        media = MediaFileUpload(file_path, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Upload and convert the file to Google Docs format
        uploaded_file = self.drive_client.files().create(body=file_metadata, media_body=media, fields="id").execute()

        # Set permissions to allow sharing
        drive_permission = {"type": "anyone", "role": "writer"}
        self.drive_client.permissions().create(fileId=uploaded_file["id"], body=drive_permission, fields="id").execute()

        # Return a shareable link
        return f"https://docs.google.com/document/d/{uploaded_file['id']}/edit"

    def authenticate(self, scopes: List[str] = settings.GCP_SCOPES):
        """
        Authenticates with Google Cloud using a service account and retrieves credentials.

        Args:
            scopes (List[str]): The required scopes for Google APIs.

        Returns:
            Credentials: Google Cloud credentials for API access.
        """
        secretmanager_client = secretmanager.SecretManagerServiceClient()
        response = secretmanager_client.access_secret_version(name=settings.GCP_SECRET_VERSION)
        service_account_info = json.loads(response.payload.data.decode("UTF-8"))
        credentials = service_account.Credentials.from_service_account_info(
            service_account_info, scopes=scopes
        )
        return credentials
